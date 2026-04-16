"""
Update 'Equation Explanation.docx' with Arrhythmia and Cardiac Failure risk logic sections.

- Inserts/updates sections at the end of the document with consistent headings and bullet lists.
- Uses python-docx for .docx manipulation.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

DOCX_NAME = 'Equation Explanation.docx'

ARRHYTHMIA_HEADING = 'Arrhythmia Risk Logic (Modified Hill Coefficients)'
CARDIAC_HEADING = 'Cardiac Failure Risk Logic (Modified Hill Coefficients)'

arrhythmia_points = [
    'Key drivers: n (dose slope), Emax (effect size), CT50_ratio (potency proxy), R² for reliability.',
    'High risk when (R² ≥ 0.5) and any of: n ≤ 1.0; or Emax ≥ 60; or CT50_ratio ≪ 1 (very potent) or ≫ 10 (marked effect at high multiples of Cmax).',
    'Typical high-risk examples: Vandetanib (n≈0.5, Emax≈36), Sunitinib (n≈0.5, Emax≈115), Epirubicin (n≈0.5, Emax≈101), Bortezomib (n≈0.57, m≈5.07).',
    'Low risk when (R² ≥ 0.5) and n ≥ 4.5 with Emax ≤ 40 and no extreme potency (CT50_ratio not extremely low).',
    'Typical low-risk examples: Chlorpromazine (n≈4.57, Emax≈39), Gemcitibine (n≈6, Emax≈23.5), Nifedipine (n≈4.98, Emax≈36.7).',
    'None/intermediate: fails both high and low criteria or R² < 0.5. Examples: Amiodarone, Dactinomycin, Etomoxir, Isoproterenol, Plicamycin.'
]

cardiac_points = [
    'Key drivers: Emax (sustained impairment), Tau (time constant), m (time steepness), n (dose slope), with R² as a reliability gate.',
    'High risk when R² ≥ 0.5 and any of: Emax ≥ 60; or (Emax ≥ 35 and Tau ≥ 70); and any of: n ≤ 2.0 or m ≥ 2.0.',
    'High-risk examples aligned with your table: Sunitinib (Emax≈115, Tau≈82, n≈0.5), Epirubicin (≈101, ≈78, 0.5), Doxorubicin (≈109, ≈86, 0.5), Erlotinib (≈73, ≈76, 0.5), Daunorubicin (≈47, ≈93, n≈1.6, m≈2.2), Sotalol (≈46, ≈72, 0.5), Vandetanib (≈36, ≈52, 0.5), Bortezomib (Emax modest but m≈5.1; included per your label).',
    'Low risk when R² ≥ 0.5 and all of: Emax ≤ 40 and Tau ≤ 60; and either n ≥ 4.5 or Emax ≤ 25.',
    'Low-risk examples: Gemcitibine (Emax≈23.5, Tau≈50, n≈6), Vincristine (≈9.9, ≈48.3). Note: Vorinostat is labeled Low in your sheet despite Emax≈115 — keep label; treat as exception.',
    'None/intermediate: fails both high and low criteria or R² < 0.5.'
]

def add_heading(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_bullets(doc: Document, items):
    # Try Word bullet styles; fall back to normal if unavailable
    bullet_styles = ['List Bullet', 'List Paragraph', None]
    for it in items:
        for style in bullet_styles:
            try:
                p = doc.add_paragraph(it, style=style) if style else doc.add_paragraph(it)
                break
            except KeyError:
                continue
        p_format = p.paragraph_format
        p_format.space_after = Pt(2)


def main():
    if not os.path.exists(DOCX_NAME):
        raise FileNotFoundError(DOCX_NAME)
    doc = Document(DOCX_NAME)

    # Append/update sections at end
    doc.add_page_break()
    add_heading(doc, ARRHYTHMIA_HEADING)
    add_bullets(doc, arrhythmia_points)

    doc.add_paragraph('')
    add_heading(doc, CARDIAC_HEADING)
    add_bullets(doc, cardiac_points)

    doc.save(DOCX_NAME)
    print(f"Updated {DOCX_NAME} with risk logic sections.")


if __name__ == '__main__':
    main()
